# -*- coding: utf-8 -*-
"""Generate a professional AI Engineer CV (.docx) for Rajveer Singh."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x0B, 0x5C, 0x8A)   # deep blue
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)
normal.font.color.rgb = DARK
pf = normal.paragraph_format
pf.space_after = Pt(2)
pf.line_spacing = 1.05

# narrow margins
for s in doc.sections:
    s.top_margin = Inches(0.5)
    s.bottom_margin = Inches(0.5)
    s.left_margin = Inches(0.7)
    s.right_margin = Inches(0.7)


def add_name(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT
    r.font.name = "Calibri"
    return p


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    r.font.color.rgb = GREY
    r.bold = True


def add_hyperlink(paragraph, url, text):
    """Insert a clickable hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # font size 9, accent color, underline
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18"); rPr.append(sz)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0B5C8A"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_contact_run(paragraph, text):
    r = paragraph.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


def add_contact():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_contact_run(p, "Jaipur, Rajasthan  •  +91 9588205114  •  ")
    add_hyperlink(p, "mailto:dev.rajveer111@gmail.com", "dev.rajveer111@gmail.com")
    add_contact_run(p, "  •  ")
    add_hyperlink(
        p,
        "https://www.linkedin.com/in/rajveer-singh-rathore-26612a234",
        "linkedin.com/in/rajveer-singh-rathore-26612a234",
    )


def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0B5C8A")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def body(text, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p


def bullet(text, sub=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    return p


def skill_line(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.05)
    rl = p.add_run(label + ": ")
    rl.bold = True
    rl.font.size = Pt(9.5)
    rl.font.color.rgb = ACCENT
    rv = p.add_run(value)
    rv.font.size = Pt(9.5)


def job_header(role, company, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    rr = p.add_run(role + "  ")
    rr.bold = True
    rr.font.size = Pt(10.5)
    rc = p.add_run("| " + company)
    rc.font.size = Pt(10)
    rc.font.color.rgb = GREY
    # tab + dates right aligned
    tabstops = p.paragraph_format.tab_stops
    tabstops.add_tab_stop(Inches(7.1), 2)  # right align
    rt = p.add_run("\t" + dates)
    rt.italic = True
    rt.font.size = Pt(9.5)
    rt.font.color.rgb = GREY


def project_header(name, stack, link=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    rn = p.add_run(name + "  ")
    rn.bold = True
    rn.font.size = Pt(10.5)
    rs = p.add_run("| " + stack)
    rs.font.size = Pt(9)
    rs.italic = True
    rs.font.color.rgb = GREY
    if link:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(1)
        url = link if link.startswith("http") else "https://" + link
        add_hyperlink(lp, url, link)


# ===================== CONTENT =====================
add_name("RAJVEER SINGH")
add_title("AI Engineer  |  Full-Stack AI Application Developer")
add_contact()

# Summary
section_heading("Summary")
body("AI Engineer and full-stack developer with 3+ years of experience building and shipping "
     "production-grade, AI-powered applications. Specialized in integrating ML and LLM systems into "
     "scalable web platforms—RAG pipelines, model-serving APIs, real-time data processing, and "
     "3D/point-cloud analytics. Combines hands-on machine-learning integration (PyTorch, Hugging Face, "
     "FastAPI) with strong end-to-end engineering across the modern web stack, enabling AI features to "
     "move from prototype to reliable, deployed product.")

# Technical Skills
section_heading("Technical Skills")
skill_line("AI / ML", "PyTorch, scikit-learn, Hugging Face Transformers, model inference & serving, "
                       "feature pipelines, model evaluation, computer vision (3D point-cloud / LiDAR)")
skill_line("LLMs & GenAI", "OpenAI & Anthropic APIs, LangChain / LlamaIndex, Retrieval-Augmented "
                            "Generation (RAG), embeddings, vector databases (pgvector, Pinecone, FAISS), "
                            "prompt engineering, AI agents")
skill_line("Languages", "Python, TypeScript, JavaScript (ES6+), SQL")
skill_line("Backend & APIs", "FastAPI, Node.js, Express.js, REST APIs, WebSockets / Socket.IO, MQTT, "
                              "authentication & authorization")
skill_line("Data", "PostgreSQL / pgvector, MongoDB, pandas, NumPy, data ingestion & ETL pipelines")
skill_line("Frontend", "React.js, Next.js, TypeScript, Tailwind CSS, Material UI, React Flow, "
                        "data visualization")
skill_line("Cloud & MLOps", "Azure App Services, Azure Blob Storage, AWS S3, AWS Cognito, Azure DevOps, "
                             "GitHub Actions / CI-CD, Docker")
skill_line("Specializations", "AI application engineering, RAG systems, real-time data platforms, "
                               "IoT & BLE analytics, computer-vision pipelines")

# Experience
section_heading("Work Experience")
job_header("AI Engineer / Full-Stack Developer", "Celebal Technologies, Jaipur", "Jul 2024 – Present")
bullet("Build AI-powered enterprise applications end to end, integrating ML/LLM models into production "
       "web platforms using Python (FastAPI), Next.js, Node.js, and TypeScript.")
bullet("Engineered an AI analytics platform for mining safety—serving deep-learning anomaly-detection "
       "models over large-scale 3D LiDAR point-cloud data with secure, scalable ingestion pipelines.")
bullet("Designed and deployed model-serving REST APIs and inference workflows, focusing on latency, "
       "throughput, security, and maintainability.")
bullet("Developed real-time data-processing and visualization features (LiDAR point clouds, BLE telemetry, "
       "live conversation analytics) for high-frequency streaming workloads.")
bullet("Collaborated with clients, product, and data teams to translate business problems into AI/ML "
       "solutions, delivering demos and driving production releases across multiple projects.")

job_header("Full-Stack Developer", "Intourist, Jaipur", "Apr 2023 – Jun 2024")
bullet("Developed and maintained web applications and travel-portal solutions using PHP, WordPress, "
       "HTML, and CSS, building responsive UIs and custom features to client requirements.")
bullet("Integrated third-party APIs and backend services to power booking workflows, content management, "
       "and customer-facing functionality.")
bullet("Contributed to performance optimization, testing, bug fixing, and production deployments across "
       "multiple client projects.")

# Projects
section_heading("Projects")

project_header("AI-Powered Mining LiDAR Analytics Platform", "Next.js, FastAPI, PyTorch, AWS S3/Cognito, pgvector")
bullet("Built a full-stack AI platform that processes, visualizes, and reports on large-scale LiDAR scans, "
       "applying deep-learning models for mining bolt-anomaly detection.")
bullet("Implemented model-serving inference APIs in FastAPI and secure, scalable file-ingestion pipelines "
       "(AWS S3 + Cognito) for high-volume point-cloud datasets.")
bullet("Developed interactive 3D point-cloud visualization and automated reporting to support engineering "
       "analysis and operational decision-making.")

project_header("RAG-Based Document Intelligence Assistant", "Python, LangChain, OpenAI/Anthropic APIs, pgvector, FastAPI")
bullet("Designed a Retrieval-Augmented Generation assistant that answers natural-language queries over "
       "private document collections using embeddings and a vector database.")
bullet("Implemented chunking, embedding, semantic retrieval, and prompt-orchestration pipelines, with "
       "source-grounded responses to reduce hallucination.")
bullet("Exposed the system through a FastAPI backend and a React chat interface with streaming responses.")

project_header("Real-Time Contact Center Intelligence Platform", "React, SignalR, TypeScript, NLP")
bullet("Built real-time agent-assist dashboards using SignalR live data streaming, surfacing call insights, "
       "sentiment indicators, and conversation metrics powered by NLP.")
bullet("Optimized frontend performance for high-frequency real-time updates while maintaining a responsive "
       "user experience.")

project_header("IoT Device Monitoring & Analytics Platform", "React, Node.js, BLE, MQTT, time-series analytics")
bullet("Developed a platform for ingesting, monitoring, and visualizing telemetry from IoT and BLE-enabled "
       "devices, with real-time dashboards and anomaly-aware reporting modules.")
bullet("Integrated BLE communication and device-configuration workflows for seamless interaction with "
       "field-deployed hardware.")

project_header("Transposer – Chrome Extension", "React, Web Audio API",
               "chromewebstore.google.com/detail/transposer-transpose-pitc/bhlokigpgcdidlakfodbghhgbdpgaobb")
bullet("Published a Chrome extension for real-time audio pitch shifting, tempo control, and equalizer "
       "adjustment using low-latency browser audio processing.")
bullet("Achieved 4,000+ weekly active users on the Chrome Web Store with a 4.5+ rating.")

# Education
section_heading("Education")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run("Master of Computer Applications (MCA)")
r.bold = True
r.font.size = Pt(10)
tabstops = p.paragraph_format.tab_stops
tabstops.add_tab_stop(Inches(7.1), 2)
rt = p.add_run("\t2022 – 2024")
rt.italic = True
rt.font.size = Pt(9.5)
rt.font.color.rgb = GREY
body("Rajasthan Technical University, Kota", space_after=2)

# Achievements
section_heading("Achievements")
bullet("Winner – Rajasthan Digifest × TiE Global Hackathon: built a Smart Canal Irrigation "
       "Monitoring & Management System using IoT sensors, LoRa, React, Node.js, MongoDB, GeoJSON, and "
       "real-time data streaming.")
bullet("Winner – AI + Data + CRM Summit Hackathon (Wipro): built an AI-driven solution focused on "
       "operational efficiency and data intelligence.")
bullet("Winner – ISIM Hackathon 2.0: created an IoT-enabled Smart Parking System (ESP32, real-time "
       "occupancy monitoring, mobile booking, web dashboards).")
bullet("First Runner-Up – HACK JKLU 3.0, a 36-hour competitive innovation hackathon.")
bullet("Third Position – IoT-Based Hackathon organized by AIET Jaipur.")

out = r"cv/ai_cv/rajveer_singh_ai_engineer.docx"
doc.save(out)
print("Saved:", out)
